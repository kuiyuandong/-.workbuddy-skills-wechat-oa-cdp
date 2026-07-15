# html_to_docx.py
# 将公众号 HTML 文章转成 .docx，供后台「文档导入」使用。
# 依赖: pip install python-docx beautifulsoup4 lxml
# 用法: python html_to_docx.py input.html output.docx
import sys
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def hex2rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def main(html_path, docx_path):
    with open(html_path, encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'Microsoft YaHei')
    style.font.size = Pt(12)

    def add_runs(p, el):
        # 处理加粗/颜色/换行
        for child in el.children:
            if child.name is None:
                txt = child.string
                if txt:
                    r = p.add_run(txt)
                    if child.parent.name in ('strong', 'b'):
                        r.bold = True
            elif child.name in ('strong', 'b'):
                for t in child.stripped_strings:
                    r = p.add_run(t); r.bold = True
            elif child.name == 'br':
                p.add_run('\n')
            elif child.name in ('span',):
                color = child.get('style', '')
                run = p.add_run(''.join(child.stripped_strings))
                if 'color:' in color:
                    hexc = color.split('color:')[1].split(';')[0].strip()
                    try: run.font.color.rgb = hex2rgb(hexc)
                    except Exception: pass

    for node in soup.body.children if soup.body else soup.children:
        if node.name is None:
            continue
        tag = node.name
        if tag in ('h1', 'h2'):
            p = doc.add_heading(level=1 if tag == 'h1' else 2)
            p.add_run(''.join(node.stripped_strings))
        elif tag == 'p':
            p = doc.add_paragraph()
            add_runs(p, node)
        elif tag in ('ul', 'ol'):
            for li in node.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                add_runs(p, li)
        elif tag == 'hr':
            doc.add_paragraph('_' * 20)
    doc.save(docx_path)
    print('✓ 已生成', docx_path)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('用法: python html_to_docx.py input.html output.docx')
        sys.exit(1)
    main(sys.argv[1], sys.argv[2])
